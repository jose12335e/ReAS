from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SOURCE = DOCS_DIR / "Manual del Sistema ReAS - contenido.md"
DOCX_OUT = DOCS_DIR / "Manual del Sistema ReAS - Reportes de Asistencia.docx"
PDF_OUT = DOCS_DIR / "Manual del Sistema ReAS - Reportes de Asistencia.pdf"
HEADER_IMAGE = ROOT / "src" / "assets" / "encabezado-jce.jpg"

BLUE = "1F4E79"
DARK = "0F172A"
GOLD = "D99A00"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, bold_prefix: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_prefix and ":" in text:
        label, rest = text.split(":", 1)
        run = paragraph.add_run(f"{label}:")
        set_run_font(run, bold=True)
        run = paragraph.add_run(rest)
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_docx_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5, bold=row_index == 0)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            set_cell_border(cell)
    doc.add_paragraph()


def parse_markdown(source: str) -> list[dict]:
    blocks: list[dict] = []
    lines = source.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            rows = []
            for table_line in table_lines:
                parts = [part.strip() for part in table_line.strip().strip("|").split("|")]
                if all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
                    continue
                rows.append(parts)
            blocks.append({"type": "table", "rows": rows})
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            blocks.append({"type": "heading", "level": level, "text": line[level:].strip()})
        elif line.startswith(">"):
            blocks.append({"type": "quote", "text": line.lstrip("> ").strip()})
        elif re.match(r"^\d+\.\s+", line):
            items = []
            while index < len(lines) and re.match(r"^\d+\.\s+", lines[index].rstrip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[index].rstrip()))
                index += 1
            blocks.append({"type": "numbered", "items": items})
            continue
        elif line.startswith("- "):
            items = []
            while index < len(lines) and lines[index].rstrip().startswith("- "):
                items.append(lines[index].rstrip()[2:])
                index += 1
            blocks.append({"type": "bullet", "items": items})
            continue
        else:
            paragraph_lines = [line]
            while index + 1 < len(lines):
                next_line = lines[index + 1].rstrip()
                if not next_line or next_line.startswith(("#", ">", "|", "- ")) or re.match(r"^\d+\.\s+", next_line):
                    break
                paragraph_lines.append(next_line)
                index += 1
            blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})
        index += 1
    return blocks


def add_docx_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.text = "Sistema ReAS | Reportes de Asistencia"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color="64748B")

    footer = section.footer.paragraphs[0]
    footer.text = "Documento de apoyo para presentación y operación interna"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=8.5, color="64748B")


def build_docx(blocks: list[dict]) -> None:
    doc = Document()
    add_docx_header_footer(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    # Cover
    if HEADER_IMAGE.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run()
        run.add_picture(str(HEADER_IMAGE), width=Inches(5.8))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sistema ReAS")
    set_run_font(run, size=26, color=BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual institucional y guía de uso para reportes de asistencia")
    set_run_font(run, size=14, color=DARK, bold=True)
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = metadata.add_run("Dirección de Gestión Humana | Unidad de Gestión de Procesos | Junio 2026")
    set_run_font(run, size=10, color="475569")
    doc.add_page_break()

    toc_heading = doc.add_paragraph("Contenido", style="Heading 1")
    sections = [block["text"] for block in blocks if block["type"] == "heading" and block["level"] in (2, 3)]
    for section_title in sections:
        add_docx_paragraph(doc, section_title, style="List Bullet")
    doc.add_page_break()

    for block in blocks:
        block_type = block["type"]
        if block_type == "heading":
            level = block["level"]
            if level == 1:
                continue
            style = "Heading 1" if level == 2 else "Heading 2" if level == 3 else "Heading 3"
            doc.add_paragraph(block["text"], style=style)
        elif block_type == "paragraph":
            add_docx_paragraph(doc, block["text"])
        elif block_type == "quote":
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_shading(cell, "FFF7E6")
            set_cell_border(cell, "EAB308")
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(block["text"])
            set_run_font(run, size=10.5, color=DARK, bold=True)
            doc.add_paragraph()
        elif block_type == "bullet":
            for item in block["items"]:
                add_docx_paragraph(doc, item, style="List Bullet", bold_prefix=True)
        elif block_type == "numbered":
            for item in block["items"]:
                add_docx_paragraph(doc, item, style="List Number", bold_prefix=True)
        elif block_type == "table":
            add_docx_table(doc, block["rows"])

    doc.core_properties.author = "Sistema ReAS"
    doc.core_properties.title = "Manual del Sistema ReAS - Reportes de Asistencia"
    doc.core_properties.subject = "Documentación institucional y guía de uso"
    doc.core_properties.comments = "Generado desde docs/Manual del Sistema ReAS - contenido.md"
    doc.save(DOCX_OUT)


def pdf_styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("ReASTitle", parent=base["Title"], fontName="Helvetica-Bold", fontSize=24, textColor=colors.HexColor(f"#{BLUE}"), alignment=TA_CENTER, spaceAfter=14),
        "subtitle": ParagraphStyle("ReASSubtitle", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=12, textColor=colors.HexColor(f"#{DARK}"), alignment=TA_CENTER, spaceAfter=20),
        "h1": ParagraphStyle("ReASH1", parent=base["Heading1"], fontName="Helvetica-Bold", fontSize=15, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=12, spaceAfter=6),
        "h2": ParagraphStyle("ReASH2", parent=base["Heading2"], fontName="Helvetica-Bold", fontSize=12, textColor=colors.HexColor(f"#{DARK}"), spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("ReASBody", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=6),
        "quote": ParagraphStyle("ReASQuote", parent=base["BodyText"], fontName="Helvetica-Bold", fontSize=9.5, leading=13, textColor=colors.HexColor(f"#{DARK}"), backColor=colors.HexColor("#FFF7E6"), borderColor=colors.HexColor("#EAB308"), borderWidth=0.75, borderPadding=8, spaceAfter=8),
        "bullet": ParagraphStyle("ReASBullet", parent=base["BodyText"], fontName="Helvetica", fontSize=9.5, leading=13, leftIndent=16, firstLineIndent=-8, spaceAfter=3),
        "small": ParagraphStyle("ReASSmall", parent=base["BodyText"], fontName="Helvetica", fontSize=8.5, leading=11, textColor=colors.HexColor("#64748B"), alignment=TA_CENTER),
    }


def table_for_pdf(rows: list[list[str]]) -> Table:
    col_count = len(rows[0])
    available_width = 6.7 * inch
    col_width = available_width / col_count
    table = Table([[Paragraph(cell, pdf_styles()["body"]) for cell in row] for row in rows], colWidths=[col_width] * col_count, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_BLUE}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{DARK}")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{BORDER}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def build_pdf(blocks: list[dict]) -> None:
    styles = pdf_styles()
    story = []
    if HEADER_IMAGE.exists():
        story.append(Image(str(HEADER_IMAGE), width=5.8 * inch, height=0.85 * inch))
        story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph("Sistema ReAS", styles["title"]))
    story.append(Paragraph("Manual institucional y guía de uso para reportes de asistencia", styles["subtitle"]))
    story.append(Paragraph("Dirección de Gestión Humana | Unidad de Gestión de Procesos | Junio 2026", styles["small"]))
    story.append(PageBreak())
    story.append(Paragraph("Contenido", styles["h1"]))
    for block in blocks:
        if block["type"] == "heading" and block["level"] in (2, 3):
            story.append(Paragraph(block["text"], styles["body"]))
    story.append(PageBreak())

    for block in blocks:
        block_type = block["type"]
        if block_type == "heading":
            if block["level"] == 1:
                continue
            story.append(Paragraph(block["text"], styles["h1"] if block["level"] == 2 else styles["h2"]))
        elif block_type == "paragraph":
            story.append(Paragraph(block["text"], styles["body"]))
        elif block_type == "quote":
            story.append(Paragraph(block["text"], styles["quote"]))
        elif block_type == "bullet":
            items = [ListItem(Paragraph(item, styles["bullet"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=12))
            story.append(Spacer(1, 4))
        elif block_type == "numbered":
            items = [ListItem(Paragraph(item, styles["bullet"])) for item in block["items"]]
            story.append(ListFlowable(items, bulletType="1", leftIndent=16))
            story.append(Spacer(1, 4))
        elif block_type == "table":
            story.append(KeepTogether([table_for_pdf(block["rows"]), Spacer(1, 8)]))

    def decorate(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#64748B"))
        canvas.drawString(0.75 * inch, 0.45 * inch, "Sistema ReAS | Reportes de Asistencia")
        canvas.drawRightString(7.75 * inch, 0.45 * inch, f"Página {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Manual del Sistema ReAS - Reportes de Asistencia",
        author="Sistema ReAS",
    )
    doc.build(story, onFirstPage=decorate, onLaterPages=decorate)


def main() -> None:
    source = SOURCE.read_text(encoding="utf-8")
    blocks = parse_markdown(source)
    build_docx(blocks)
    build_pdf(blocks)
    print(f"Generated: {DOCX_OUT}")
    print(f"Generated: {PDF_OUT}")


if __name__ == "__main__":
    main()
